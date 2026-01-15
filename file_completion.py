# file_completion.py

import os
import uuid
import google.generativeai as genai
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import docx.oxml.shared
import docx.text.paragraph
import docx.enum.style
from werkzeug.utils import secure_filename
import json
import base64 # Import base64
from apikey import GEN_AI_API_KEY  # Ensure you have your API key set up in apikeys.py

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "www/uploaded_files"
COMPLETED_FOLDER = "www/completed_files"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(COMPLETED_FOLDER, exist_ok=True)

# Replace with your actual API key
genai.configure(api_key=GEN_AI_API_KEY)
model = genai.GenerativeModel("gemini-1.5-flash")

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def insert_paragraph_after(doc_obj, paragraph_before, text_to_insert, style_name=None):
    """
    Inserts a new paragraph with 'text_to_insert' after 'paragraph_before'.
    'doc_obj' is the Document object.
    'style_name' is the name of a PARAGRAPH style (e.g., 'Normal', 'Heading 1').
    """
    if paragraph_before.runs:
        parent = paragraph_before._element.getparent()
        # Create a new paragraph element
        new_p = docx.oxml.shared.OxmlElement("w:p")
        # Add the new paragraph after the target paragraph
        paragraph_before._element.addnext(new_p)
        # Wrap the new paragraph element in a Paragraph object
        new_paragraph = docx.text.paragraph.Paragraph(new_p, paragraph_before.document)
        new_paragraph.add_run(text_to_insert)
        if style_name:
            new_paragraph.style = doc_obj.styles[style_name]
        return new_paragraph
    else:
        # If the paragraph is empty, just append to the document
        return doc_obj.add_paragraph(text_to_insert, style=style_name)

# --- New Test Routes ---
@app.route("/", methods=["GET"])
def home():
    """A simple root route to confirm the Flask service is running."""
    return "Flask file completion service is running!", 200

@app.route("/test", methods=["GET"])
def test_route():
    """A test route to verify basic API connectivity."""
    return jsonify({"message": "Test route successful!", "status": "ok"}), 200

@app.route("/complete-file", methods=["POST"])
def complete_file_endpoint():
    print("DEBUG: /complete-file endpoint hit.")
    if not request.json:
        print("ERROR: Request must be JSON.")
        return jsonify({"error": "Request must be JSON"}), 400

    filename = request.json.get("filename")
    base64_file_data = request.json.get("file_data")

    if not filename or not base64_file_data:
        print("ERROR: Missing filename or file_data in request.")
        return jsonify({"error": "Missing filename or file_data"}), 400

    temp_uploaded_filepath = None
    try:
        # Decode the base64 data and save it as a temporary file
        file_bytes = base64.b64decode(base64_file_data)
        unique_filename = f"{uuid.uuid4()}_{secure_filename(filename)}"
        temp_uploaded_filepath = os.path.join(UPLOAD_FOLDER, unique_filename)

        with open(temp_uploaded_filepath, "wb") as f:
            f.write(file_bytes)
        print(f"DEBUG: Temporary file saved at: {temp_uploaded_filepath}")

        extracted_text = extract_text_from_docx(temp_uploaded_filepath)
        print("DEBUG: Extracted text from DOCX.")

        prompt = f"You are a helpful assistant. Complete the following document based on its content. Provide a detailed comprehensive and professional completion. The document content is:\n\n{extracted_text}"
        print(f"DEBUG: Sending prompt to Generative Model. Prompt length: {len(prompt)}")

        response = model.generate_content(prompt)
        gemini_response_text = response.text
        print("DEBUG: Received response from Generative Model.")

        # Determine the output filename
        name, ext = os.path.splitext(filename)
        completed_filename = f"{name}_completed{ext}"
        completed_file_path = os.path.join(COMPLETED_FOLDER, completed_filename)

        # Process the DOCX for targeted insertion or append
        doc = Document(temp_uploaded_filepath)

        # Attempt to parse Gemini's response as JSON for structured insertion
        try:
            structured_content = json.loads(gemini_response_text)
            if isinstance(structured_content, dict):
                for key, value in structured_content.items():
                    # Find the last paragraph in the document
                    last_paragraph = doc.paragraphs[-1] if doc.paragraphs else None

                    key_paragraph = doc.add_paragraph(f"{key}:", style='Normal')
                    try:
                        # Try to apply 'Heading 3' style for the key
                        key_paragraph.style = doc.styles['Heading 3']
                        for run in key_paragraph.runs:
                            run.bold = True
                    except KeyError:
                        # Fallback to 'Normal' style if 'Heading 3' doesn't exist
                        key_paragraph.style = doc.styles['Normal']
                        for run in key_paragraph.runs:
                            run.bold = True
                    except Exception as e:
                        print(f"WARNING: Could not apply 'Heading 3' or 'Normal' with bolding for key '{key}': {e}. Adding as plain normal paragraph.")
                        key_paragraph = doc.add_paragraph(f"{key}:", style='Normal')

                    for line in str(value).split("\n"):
                        if line.strip():
                            doc.add_paragraph(line, style='Normal')
            else:
                # If not a dictionary, treat as plain text and append
                print("DEBUG: Gemini response is not a dictionary. Appending as plain text.")
                for line in gemini_response_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line, style='Normal')

        except json.JSONDecodeError:
            print("DEBUG: Gemini response is not valid JSON. Appending as plain text.")
            # If Gemini's response is not JSON, just append it as paragraphs
            for line in gemini_response_text.split("\n"):
                if line.strip(): # Only add non-empty lines
                    doc.add_paragraph(line, style='Normal')

        doc.save(completed_file_path)
        print(f"DEBUG: Completed file saved at: {completed_file_path}")

        # Read the completed file as base64 and send it back
        with open(completed_file_path, "rb") as f:
            completed_file_base64_encoded = base64.b64encode(f.read()).decode('utf-8')

        return jsonify({
            "completed_filename": completed_filename,
            "completed_file_data": completed_file_base64_encoded
        })

    except Exception as e:
        print(f"ERROR: An error occurred during file completion: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        if temp_uploaded_filepath and os.path.exists(temp_uploaded_filepath):
            os.remove(temp_uploaded_filepath)
            print(f"DEBUG: Cleaned up temporary uploaded file: {temp_uploaded_filepath}")

# This block is for direct execution of file_completion.py for testing.
# In your setup, main.py imports and runs this app in a thread, so this __name__ == "__main__"
# block won't execute when run via main.py.
if __name__ == "__main__":
    app.run(port=5002, debug=True) # Run with debug for more verbose error messages
