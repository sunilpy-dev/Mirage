# worksheet.py
import os
import uuid
import base64
import json
import PyPDF2
import pytesseract
import google.generativeai as genai
from flask import Flask, request, jsonify # Removed send_file as it will be replaced
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from werkzeug.utils import secure_filename
from PIL import Image

# --- API Key Configuration ---
# IMPORTANT: Replace "" with your actual Gemini API key.
# If you are using an apikey.py file, ensure GEN_AI_API_KEY is correctly defined there.
try:
    from apikey import GEN_AI_API_KEY
    if not GEN_AI_API_KEY:
        raise ValueError("GEN_AI_API_KEY is empty in apikey.py. Please provide your API key.")
except ImportError:
    print("WARNING: apikey.py not found. Please set your API key directly in this file.")
    GEN_AI_API_KEY = os.getenv("GEMINI_API_KEY", "")

    if not GEN_AI_API_KEY:
        print("CRITICAL ERROR: GEMINI_API_KEY is not set. Please set it in apikey.py or as an environment variable.")

# --- Tesseract OCR Configuration ---
# IMPORTANT: You must install Tesseract OCR on your system.
# Download from: https://tesseract-ocr.github.io/tessdoc/Installation.html
# After installation, set the path to your tesseract executable if it's not in your system's PATH.
# Example for Windows: pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# Example for Linux/macOS: pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'
try:
    # Attempt to set a common default path or rely on system PATH
    # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' # Uncomment and modify for your system
    pass
except Exception as e:
    print(f"WARNING: pytesseract configuration failed. OCR functionality may be limited. Error: {e}")

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "www/uploaded_files"
WORKSHEET_FOLDER = "www/worksheets"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(WORKSHEET_FOLDER, exist_ok=True)

# Configure Gemini AI model
if GEN_AI_API_KEY:
    try:
        genai.configure(api_key=GEN_AI_API_KEY)
        model = genai.GenerativeModel("gemini-2.5-flash")
    except Exception as e:
        print(f"CRITICAL ERROR: Failed to configure Gemini API with provided key: {e}")
        model = None
else:
    print("CRITICAL ERROR: Gemini API key is missing. Model will not be available.")
    model = None

def extract_text_from_docx(file_path):
    """
    Extracts all text from a .docx file.
    """
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"ERROR: Could not extract text from DOCX {file_path}: {e}")
        return ""

def extract_text_from_image(image_path):
    """
    Extracts text from an image file using OCR (pytesseract).
    """
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return text
    except pytesseract.TesseractNotFoundError:
        print("ERROR: Tesseract OCR is not installed or not found in PATH. Image text extraction will fail.")
        return ""
    except Exception as e:
        print(f"ERROR: Could not perform OCR on image {image_path}: {e}")
        return ""

def extract_text_from_pdf(file_path):
    """
    Extracts all text from a PDF file.
    """
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"ERROR: Could not extract text from PDF {file_path}: {e}")
        return ""

def generate_questions_with_gemini(extracted_text, mime_type, base64_file_data, target_level=None):
    """
    Generates questions using the Gemini API based on the extracted text or image data,
    requesting multiple question types and difficulty levels.
    An optional target_level can be provided to request questions for a specific level only.
    """
    if model is None:
        raise Exception("Gemini API model is not configured.")

    gemini_input_parts = []
    
    if mime_type.startswith("image/"):
        gemini_input_parts.append({
            "inline_data": {
                "mime_type": mime_type,
                "data": base64_file_data
            }
        })
        gemini_input_parts.append({
            "text": f"Generate questions based on the content of this image. The extracted text from the image is:\n\n{extracted_text if extracted_text else 'No legible text was extracted by OCR. Please analyze the image content directly.'}"
        })
    else:
        gemini_input_parts.append({"text": extracted_text})

    level_specific_prompt = ""
    if target_level:
        level_map = {
            1: "Multiple Choice Questions (MCQ). Each MCQ must have exactly 4 options.",
            2: "Fill in the Blanks questions. Use '[BLANK]' to indicate where the blank is.",
            3: "Short Answer questions.",
            4: "Long Answer questions."
        }
        level_type = level_map.get(target_level, "various question types")
        level_specific_prompt = f"Focus ONLY on Level {target_level}: {level_type}. "
        
        # Adjust JSON structure prompt based on target_level
        if target_level == 1:
            level_specific_prompt += """
            Provide the questions in a JSON array format. Each item in the array should be an object with the following structure:
            {
                "level": 1,
                "question": "The question text",
                "type": "Multiple Choice",
                "difficulty": "e.g., Easy, Medium, Hard",
                "topic": "The topic from the document the question relates to (if identifiable)",
                "options": ["Option A", "Option B", "Option C", "Option D"],
                "correct_answer": "The correct option (e.g., 'Option A' or the full text of the correct option)"
            }
            """
        elif target_level == 2:
            level_specific_prompt += """
            Provide the questions in a JSON array format. Each item in the array should be an object with the following structure:
            {
                "level": 2,
                "question": "The sentence with [BLANK] placeholder(s).",
                "type": "Fill in the Blanks",
                "difficulty": "e.g., Easy, Medium, Hard",
                "topic": "The topic from the document the question relates to (if identifiable)",
                "correct_answer": ["Correct word 1", "Correct word 2"] // Array for multiple blanks, string for single
            }
            """
        elif target_level == 3:
            level_specific_prompt += """
            Provide the questions in a JSON array format. Each item in the array should be an object with the following structure:
            {
                "level": 3,
                "question": "The short answer question text.",
                "type": "Short Answer",
                "difficulty": "e.g., Easy, Medium, Hard",
                "topic": "The topic from the document the question relates to (if identifiable)",
                "suggested_answer": "A concise suggested answer."
            }
            """
        elif target_level == 4:
            level_specific_prompt += """
            Provide the questions in a JSON array format. Each item in the array should be an object with the following structure:
            {
                "level": 4,
                "question": "The long answer question text.",
                "type": "Long Answer",
                "difficulty": "e.g., Easy, Medium, Hard",
                "topic": "The topic from the document the question relates to (if identifiable)",
                "suggested_answer": "A comprehensive suggested answer."
            }
            """
        level_specific_prompt += " Please generate at least 10 questions for this level."
    else:
        level_specific_prompt = """
        Generate questions across four levels with specific types for each:
        - Level 1: Multiple Choice Questions (MCQ). Each MCQ must have exactly 4 options.
        - Level 2: Fill in the Blanks questions. Use '[BLANK]' to indicate where the blank is.
        - Level 3: Short Answer questions.
        - Level 4: Long Answer questions.

        Provide the questions in a JSON array format. Each item in the array should be an object with the following structure,
        adapting to the question type:

        For Multiple Choice Questions (Level 1):
        {
            "level": 1,
            "question": "The question text",
            "type": "Multiple Choice",
            "difficulty": "e.g., Easy, Medium, Hard",
            "topic": "The topic from the document the question relates to (if identifiable)",
            "options": ["Option A", "Option B", "Option C", "Option D"],
            "correct_answer": "The correct option (e.g., 'Option A' or the full text of the correct option)"
        }

        For Fill in the Blanks Questions (Level 2):
        {
            "level": 2,
            "question": "The sentence with [BLANK] placeholder(s).",
            "type": "Fill in the Blanks",
            "difficulty": "e.g., Easy, Medium, Hard",
            "topic": "The topic from the document the question relates to (if identifiable)",
            "correct_answer": ["Correct word 1", "Correct word 2"] // Array for multiple blanks, string for single
        }

        For Short Answer Questions (Level 3):
        {
            "level": 3,
            "question": "The short answer question text.",
            "type": "Short Answer",
            "difficulty": "e.g., Easy, Medium, Hard",
            "topic": "The topic from the document the question relates to (if identifiable)",
            "suggested_answer": "A concise suggested answer."
        }

        For Long Answer Questions (Level 4):
        {
            "level": 4,
            "question": "The long answer question text.",
            "type": "Long Answer",
            "difficulty": "e.g., Easy, Medium, Hard",
            "topic": "The topic from the document the question relates to (if identifiable)",
            "suggested_answer": "A comprehensive suggested answer."
        }

        Please generate at least 10 questions for each level.
        """

    final_prompt = "You are a question generation bot. Your task is to create a set of diverse and challenging questions based on the provided document content. The questions should cover various topics and assess understanding at different levels. " + level_specific_prompt

    if mime_type.startswith("image/"):
        gemini_input_parts[1]["text"] += "\n\n" + final_prompt
    else:
        gemini_input_parts.append({"text": final_prompt})

    try:
        response = model.generate_content(gemini_input_parts)
        gemini_response_text = response.text

        # Check for empty Gemini response before parsing
        if not gemini_response_text or not gemini_response_text.strip():
            # This is a critical point. If Gemini returns nothing, it's an issue.
            # Raise a specific error that can be caught and returned as JSON by the Flask endpoint.
            raise ValueError("Gemini API returned an empty or whitespace-only response. This usually means the model couldn't generate content for the given prompt.")

        # Clean up potential markdown formatting from Gemini's JSON response
        if gemini_response_text.startswith("```json"):
            gemini_response_text = gemini_response_text[7:].strip()
        if gemini_response_text.endswith("```"):
            gemini_response_text = gemini_response_text[:-3].strip()

        # Attempt to parse the JSON. If this fails, the JSON is malformed.
        try:
            questions_data = json.loads(gemini_response_text)
        except json.JSONDecodeError as e:
            # If JSON parsing fails, it means Gemini's output was not valid JSON.
            # Include the problematic text in the error for debugging.
            raise ValueError(f"Gemini API response was not valid JSON. Error: {e}. Raw response: '{gemini_response_text[:500]}...'")

        if not isinstance(questions_data, list):
            raise ValueError("Expected a JSON array of questions from Gemini.")

        # Basic validation for each question object based on its type
        for q in questions_data:
            if not all(k in q for k in ["level", "question", "type", "difficulty"]):
                raise ValueError(f"Each question object must contain 'level', 'question', 'type', and 'difficulty' fields. Missing in: {q}")
            
            if q["type"] == "Multiple Choice":
                if "options" not in q or not isinstance(q["options"], list) or len(q["options"]) != 4:
                    raise ValueError(f"MCQ type questions must have an 'options' array with 4 items. Issue in: {q}")
                if "correct_answer" not in q or not isinstance(q["correct_answer"], str):
                    raise ValueError(f"MCQ type questions must have a 'correct_answer' string. Issue in: {q}")
            elif q["type"] == "Fill in the Blanks":
                if "correct_answer" not in q or not (isinstance(q["correct_answer"], str) or isinstance(q["correct_answer"], list)):
                    raise ValueError(f"Fill in the Blanks questions must have a 'correct_answer' (string or list). Issue in: {q}")
            elif q["type"] in ["Short Answer", "Long Answer"]:
                if "suggested_answer" not in q or not isinstance(q["suggested_answer"], str):
                    raise ValueError(f"{q['type']} questions must have a 'suggested_answer' string. Issue in: {q}")
            else:
                raise ValueError(f"Unsupported question type: {q['type']}. Issue in: {q}")
        
        return questions_data
    except Exception as e:
        # Catch any exception during Gemini interaction or parsing here,
        # and re-raise it with context. The endpoint will then catch this.
        raise Exception(f"Error during Gemini question generation or parsing: {e}")

def create_worksheet_docx(questions_data, filename="worksheet.docx", target_level=None):
    """
    Creates a DOCX file from the generated questions, categorized by levels and types.
    If target_level is provided, only questions for that level will be included.
    """
    document = Document()

    # Set default font size for the document
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    document.add_heading("Generated Worksheet", level=0) # Main Title
    document.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Group questions by level
    questions_by_level = {1: [], 2: [], 3: [], 4: []}
    for q in questions_data:
        level = q.get("level")
        if level in questions_by_level:
            questions_by_level[level].append(q)
        else:
            print(f"WARNING: Question with unsupported level {level} found: {q.get('question', 'N/A')}")

    # Add questions for each level
    levels_to_include = [target_level] if target_level else sorted(questions_by_level.keys())

    for level_num in levels_to_include:
        level_questions = questions_by_level[level_num]
        if not level_questions:
            continue

        level_title = ""
        if level_num == 1:
            level_title = "Level 1: Multiple Choice Questions"
        elif level_num == 2:
            level_title = "Level 2: Fill in the Blanks"
        elif level_num == 3:
            level_title = "Level 3: Short Answer Questions"
        elif level_num == 4:
            level_title = "Level 4: Long Answer Questions"
        
        document.add_heading(level_title, level=1)
        document.add_paragraph("\n") # Add space after level heading

        for i, q in enumerate(level_questions):
            p = document.add_paragraph()
            p.add_run(f"{i+1}. {q.get('question', 'N/A')}\n").bold = True
            
            # Add topic and difficulty
            details_run = p.add_run(f"   (Difficulty: {q.get('difficulty', 'N/A')}")
            if q.get('topic'):
                details_run.add_text(f", Topic: {q.get('topic', 'N/A')})")
            else:
                details_run.add_text(")")
            details_run.font.size = Pt(10)
            details_run.font.italic = True
            
            if q.get("type") == "Multiple Choice" and q.get("options"):
                for j, option in enumerate(q["options"]):
                    document.add_paragraph(f"  {chr(65 + j)}. {option}", style='List Bullet')
                document.add_paragraph("\n") # Add space after options
            elif q.get("type") == "Fill in the Blanks":
                document.add_paragraph("  ____________________________________________________________________\n")
                document.add_paragraph("  ____________________________________________________________________\n")
            elif q.get("type") == "Short Answer":
                document.add_paragraph("  ____________________________________________________________________\n")
                document.add_paragraph("  ____________________________________________________________________\n")
                document.add_paragraph("  ____________________________________________________________________\n")
            elif q.get("type") == "Long Answer":
                for _ in range(8): # Provide more lines for long answers
                    document.add_paragraph("  ____________________________________________________________________\n")
            
            document.add_paragraph("\n") # Add some space between questions

    # Add Answer Key
    document.add_page_break()
    document.add_heading("Answer Key", level=0)
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for level_num in levels_to_include:
        level_questions = questions_by_level[level_num]
        if not level_questions:
            continue
        
        level_title = ""
        if level_num == 1:
            level_title = "Level 1: Multiple Choice Answers"
        elif level_num == 2:
            level_title = "Level 2: Fill in the Blanks Answers"
        elif level_num == 3:
            level_title = "Level 3: Short Answer Suggestions"
        elif level_num == 4:
            level_title = "Level 4: Long Answer Suggestions"
        
        document.add_heading(level_title, level=2)

        for i, q in enumerate(level_questions):
            p = document.add_paragraph()
            p.add_run(f"Q{i+1}: {q.get('question', 'N/A')}\n").bold = True
            
            if q.get("type") == "Multiple Choice" and q.get("correct_answer"):
                p.add_run(f"   Correct Answer: {q['correct_answer']}")
            elif q.get("type") == "Fill in the Blanks" and q.get("correct_answer"):
                answer_text = q['correct_answer']
                if isinstance(answer_text, list):
                    answer_text = ", ".join(answer_text)
                p.add_run(f"   Answer: {answer_text}")
            elif q.get("type") in ["Short Answer", "Long Answer"] and q.get("suggested_answer"):
                p.add_run(f"   Suggested Answer: {q['suggested_answer']}")
            
            document.add_paragraph("\n") # Add space

    file_path = os.path.join(WORKSHEET_FOLDER, filename)
    document.save(file_path)
    return file_path

@app.route("/", methods=["GET"])
def home():
    """A simple root route to confirm the Flask service is running."""
    return "Flask Worksheet Generator service is running!", 200

@app.route("/generate-worksheet", methods=["POST"])
def generate_worksheet_endpoint():
    print("DEBUG: /generate-worksheet endpoint hit.")
    if not request.json:
        print("ERROR: Request must be JSON.")
        return jsonify({"error": "Request must be JSON"}), 400

    filename = request.json.get("filename")
    base64_file_data = request.json.get("file_data")
    mime_type = request.json.get("mime_type")
    target_level = request.json.get("level", None)

    if not filename or not base64_file_data or not mime_type:
        print("ERROR: Missing filename, file_data, or mime_type in request.")
        return jsonify({"error": "Missing filename, file_data, or mime_type"}), 400

    if model is None:
        print("ERROR: Gemini API model is not configured.")
        return jsonify({"error": "Gemini API model is not configured. Please check your API key."}), 500

    temp_uploaded_filepath = None
    worksheet_filepath = None # Initialize to None for finally block
    try:
        file_bytes = base64.b64decode(base64_file_data)
        unique_filename = f"{uuid.uuid4()}_{secure_filename(filename)}"
        temp_uploaded_filepath = os.path.join(UPLOAD_FOLDER, unique_filename)

        with open(temp_uploaded_filepath, "wb") as f:
            f.write(file_bytes)
        print(f"DEBUG: Temporary file saved at: {temp_uploaded_filepath}")

        extracted_text = ""
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extracted_text = extract_text_from_docx(temp_uploaded_filepath)
            print("DEBUG: Extracted text from DOCX.")
        elif mime_type == "application/pdf":
            extracted_text = extract_text_from_pdf(temp_uploaded_filepath)
            print("DEBUG: Extracted text from PDF.")
        elif mime_type.startswith("image/"):
            extracted_text = extract_text_from_image(temp_uploaded_filepath)
            print(f"DEBUG: Processed image and extracted text (length: {len(extracted_text)}).")
        else:
            print("ERROR: Unsupported file type. Only .docx, .pdf, .png, .jpg, .jpeg are supported.")
            return jsonify({"error": "Unsupported file type. Only .docx, .pdf, .png, .jpg, .jpeg are supported."}), 400
        
        if not extracted_text and not mime_type.startswith("image/"):
            print("WARNING: No text extracted from the document/PDF. This might result in poor question generation.")
            # Still proceed, but log warning

        # Generate questions using Gemini, passing the target_level
        try:
            questions_data = generate_questions_with_gemini(extracted_text, mime_type, base64_file_data, target_level)
            print(f"DEBUG: Generated {len(questions_data)} questions.")
        except ValueError as gemini_val_exc: # Catch specific ValueError from generate_questions_with_gemini
            print(f"ERROR: Gemini question generation/parsing issue: {gemini_val_exc}")
            return jsonify({"error": f"Gemini content generation/parsing failed: {gemini_val_exc}"}), 500
        except Exception as gemini_exc: # Catch other general exceptions
            print(f"ERROR: General error during Gemini question generation: {gemini_exc}")
            return jsonify({"error": f"An unexpected error occurred during AI question generation: {gemini_exc}"}), 500

        # Create DOCX worksheet, passing the target_level for filtering
        worksheet_filename = f"worksheet_{str(target_level) if target_level else 'all'}_level_{uuid.uuid4()}.docx"
        worksheet_filepath = create_worksheet_docx(questions_data, worksheet_filename, target_level)
        print(f"DEBUG: Worksheet DOCX created at: {worksheet_filepath}")

        # --- MODIFIED: Send the DOCX file as base64 in a JSON response ---
        try:
            with open(worksheet_filepath, "rb") as f:
                completed_file_base64_encoded = base64.b64encode(f.read()).decode('utf-8')

            return jsonify({
                "completed_filename": worksheet_filename,
                "completed_file_data": completed_file_base64_encoded
            })
        except Exception as encode_send_err:
            print(f"ERROR: Failed to encode and send worksheet file: {encode_send_err}")
            return jsonify({"error": f"Failed to encode and send worksheet file: {encode_send_err}"}), 500

    except Exception as e: # Catch any other unexpected errors during file processing
        print(f"ERROR: An error occurred during worksheet generation: {e}")
        # Ensure that any unhandled exception also returns a JSON response
        return jsonify({"error": f"An unexpected error occurred: {str(e)}"}), 500
    finally:
        # Clean up temporary uploaded file
        if temp_uploaded_filepath and os.path.exists(temp_uploaded_filepath):
            os.remove(temp_uploaded_filepath)
            print(f"DEBUG: Cleaned up temporary uploaded file: {temp_uploaded_filepath}")
        # Clean up generated worksheet file after sending (or attempting to send)
        if worksheet_filepath and os.path.exists(worksheet_filepath):
            os.remove(worksheet_filepath)
            print(f"DEBUG: Cleaned up generated worksheet file: {worksheet_filepath}")


@app.errorhandler(Exception)
def handle_exception(e):
    """Global error handler to ensure all errors return JSON."""
    import traceback
    print(f"ERROR: Uncaught exception: {e}\n{traceback.format_exc()}")
    return jsonify({"error": f"Internal server error: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(port=5008, debug=True)