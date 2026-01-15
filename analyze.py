# analyze.py

import os
import uuid
import google.generativeai as genai
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import docx.oxml.shared
import docx.text.paragraph
import docx.enum.style
from werkzeug.utils import secure_filename
import json
import base64
from apikey import GEN_AI_API_KEY  # Ensure you have your API key set up in apikeys.py

app = Flask(__name__)
CORS(app)

# Define distinct folders for analyze.py to avoid conflicts with file_completion.py
UPLOAD_FOLDER = "www/analyze_uploaded_files"
COMPLETED_FOLDER = "www/analyze_completed_files"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(COMPLETED_FOLDER, exist_ok=True)

# Configure the Generative AI model
genai.configure(api_key=GEN_AI_API_KEY)
model = genai.GenerativeModel("gemini-2.5-flash") # Using gemini-2.5-flash as requested implicitly

def extract_text_from_docx(file_path):
    """
    Extracts all text from a .docx file.
    """
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def insert_paragraph_after(doc_obj, paragraph_before, text_to_insert, style_name=None):
    """
    Inserts a new paragraph with 'text_to_insert' after 'paragraph_before'.
    'doc_obj' is the Document object.
    'style_name' is the name of a PARAGRAPH style (e.g., 'Normal', 'Heading 1').
    This function handles inserting after existing paragraphs.
    """
    parent = paragraph_before._element.getparent()
    new_p = docx.oxml.shared.OxmlElement("w:p")
    paragraph_before._element.addnext(new_p)
    new_paragraph = docx.text.paragraph.Paragraph(new_p, paragraph_before.document)
    new_paragraph.add_run(text_to_insert)
    if style_name:
        try:
            new_paragraph.style = doc_obj.styles[style_name]
        except KeyError:
            print(f"WARNING: Style '{style_name}' not found. Using 'Normal' style.")
            new_paragraph.style = doc_obj.styles['Normal']
    return new_paragraph

# --- Test Routes ---
@app.route("/", methods=["GET"])
def home():
    """A simple root route to confirm the Flask service is running."""
    return "Flask file analysis service is running!", 200

@app.route("/test", methods=["GET"])
def test_route():
    """A test route to verify basic API connectivity."""
    return jsonify({"message": "Test route successful for analyze.py!", "status": "ok"}), 200

@app.route("/analyze-file", methods=["POST"])
def analyze_file_endpoint():
    """
    Endpoint to receive a DOCX file, extract its text, send it to the Generative Model
    for scientific completion/analysis, and return the completed DOCX as base64.
    """
    print("DEBUG: /analyze-file endpoint hit.")
    if not request.json:
        print("ERROR: Request must be JSON.")
        return jsonify({"error": "Request must be JSON"}), 400

    filename = request.json.get("filename")
    base64_file_data = request.json.get("file_data")

    if not filename or not base64_file_data:
        print("ERROR: Missing filename or file_data in request.")
        return jsonify({"error": "Missing filename or file_data"}), 400

    temp_uploaded_filepath = None
    try:
        # Decode the base64 data and save it as a temporary file
        file_bytes = base64.b64decode(base64_file_data)
        unique_filename = f"{uuid.uuid4()}_{secure_filename(filename)}"
        temp_uploaded_filepath = os.path.join(UPLOAD_FOLDER, unique_filename)

        with open(temp_uploaded_filepath, "wb") as f:
            f.write(file_bytes)
        print(f"DEBUG: Temporary file saved at: {temp_uploaded_filepath}")

        extracted_text = extract_text_from_docx(temp_uploaded_filepath)
        print("DEBUG: Extracted text from DOCX.")

        # Modified prompt for scientific and detailed analysis
        prompt = (f"You are an expert scientific researcher. Analyze the following document content and "
                  f"provide a very much extensively detailed, comprehensive, and scientifically rigorous completion or analysis. "
                  f"Use precise scientific language, include relevant technical terms, and cite potential "
                  f"areas for further research or consideration if applicable. "
                  f"The document content is:\n\n{extracted_text}")
        print(f"DEBUG: Sending prompt to Generative Model for scientific analysis. Prompt length: {len(prompt)}")

        response = model.generate_content(prompt)
        gemini_response_text = response.text
        print("DEBUG: Received response from Generative Model.")

        # Determine the output filename
        name, ext = os.path.splitext(filename)
        completed_filename = f"{name}_analyzed{ext}"
        completed_file_path = os.path.join(COMPLETED_FOLDER, completed_filename)

        # Process the DOCX for insertion
        doc = Document(temp_uploaded_filepath)

        # Attempt to parse Gemini's response as JSON for structured insertion
        try:
            structured_content = json.loads(gemini_response_text)
            if isinstance(structured_content, dict):
                for key, value in structured_content.items():
                    # Add heading for the key
                    key_paragraph = doc.add_paragraph(f"{key}:", style='Normal')
                    try:
                        # Try to apply 'Heading 3' style for the key and bold it
                        key_paragraph.style = doc.styles['Heading 3']
                        for run in key_paragraph.runs:
                            run.bold = True
                    except KeyError:
                        # Fallback to 'Normal' style if 'Heading 3' doesn't exist, and bold
                        print(f"WARNING: Style 'Heading 3' not found. Using 'Normal' for '{key}'.")
                        key_paragraph.style = doc.styles['Normal']
                        for run in key_paragraph.runs:
                            run.bold = True
                    except Exception as e:
                        print(f"WARNING: Could not apply bolding/style for key '{key}': {e}. Adding as plain normal paragraph.")
                        key_paragraph = doc.add_paragraph(f"{key}:", style='Normal') # Add as plain if all else fails

                    # Add the value content as normal paragraphs
                    for line in str(value).split("\n"):
                        if line.strip():
                            doc.add_paragraph(line, style='Normal')
            else:
                # If not a dictionary, treat as plain text and append
                print("DEBUG: Gemini response is not a dictionary. Appending as plain text.")
                for line in gemini_response_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line, style='Normal')

        except json.JSONDecodeError:
            print("DEBUG: Gemini response is not valid JSON. Appending as plain text.")
            # If Gemini's response is not JSON, just append it as paragraphs
            for line in gemini_response_text.split("\n"):
                if line.strip(): # Only add non-empty lines
                    doc.add_paragraph(line, style='Normal')

        doc.save(completed_file_path)
        print(f"DEBUG: Completed file saved at: {completed_file_path}")

        # Read the completed file as base64 and send it back
        with open(completed_file_path, "rb") as f:
            completed_file_base64_encoded = base64.b64encode(f.read()).decode('utf-8')

        return jsonify({
            "completed_filename": completed_filename,
            "completed_file_data": completed_file_base64_encoded
        })

    except Exception as e:
        print(f"ERROR: An error occurred during file analysis: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        if temp_uploaded_filepath and os.path.exists(temp_uploaded_filepath):
            os.remove(temp_uploaded_filepath)
            print(f"DEBUG: Cleaned up temporary uploaded file: {temp_uploaded_filepath}")

@app.route("/summarize-file", methods=["POST"])
def summarize_file_endpoint():
    """
    Endpoint to receive a DOCX file, extract its text, send it to the Generative Model
    for summarization, and return the summarized DOCX as base64.
    """
    print("DEBUG: /summarize-file endpoint hit.")
    if not request.json:
        print("ERROR: Request must be JSON.")
        return jsonify({"error": "Request must be JSON"}), 400

    filename = request.json.get("filename")
    base64_file_data = request.json.get("file_data")

    if not filename or not base64_file_data:
        print("ERROR: Missing filename or file_data in request.")
        return jsonify({"error": "Missing filename or file_data"}), 400

    temp_uploaded_filepath = None
    try:
        # Decode the base64 data and save it as a temporary file
        file_bytes = base64.b64decode(base64_file_data)
        unique_filename = f"{uuid.uuid4()}_{secure_filename(filename)}"
        temp_uploaded_filepath = os.path.join(UPLOAD_FOLDER, unique_filename)

        with open(temp_uploaded_filepath, "wb") as f:
            f.write(file_bytes)
        print(f"DEBUG: Temporary file saved at: {temp_uploaded_filepath}")

        extracted_text = extract_text_from_docx(temp_uploaded_filepath)
        print("DEBUG: Extracted text from DOCX.")

        # Prompt for summarization
        prompt = (f"Summarize the following document content concisely and accurately. "
                  f"Provide the summary in clear, readable paragraphs. "
                  f"Document content:\n\n{extracted_text}")
        print(f"DEBUG: Sending prompt to Generative Model for summarization. Prompt length: {len(prompt)}")

        response = model.generate_content(prompt)
        gemini_response_text = response.text
        print("DEBUG: Received response from Generative Model.")

        # Determine the output filename
        name, ext = os.path.splitext(filename)
        completed_filename = f"{name}_summarized{ext}"
        completed_file_path = os.path.join(COMPLETED_FOLDER, completed_filename)

        # Process the DOCX for insertion
        doc = Document(temp_uploaded_filepath)
        
        # Add a heading for the summary
        doc.add_paragraph("--- Summary ---", style='Normal')
        try:
            summary_heading = doc.paragraphs[-1] # Get the last added paragraph
            summary_heading.style = doc.styles['Heading 2']
            for run in summary_heading.runs:
                run.bold = True
        except KeyError:
            print("WARNING: Style 'Heading 2' not found. Using 'Normal' for summary heading.")
            for run in doc.paragraphs[-1].runs:
                run.bold = True
        except Exception as e:
            print(f"WARNING: Could not apply style for summary heading: {e}")
            
        # Append the summary text
        for line in gemini_response_text.split("\n"):
            if line.strip():
                doc.add_paragraph(line, style='Normal')

        doc.save(completed_file_path)
        print(f"DEBUG: Completed file saved at: {completed_file_path}")

        # Read the completed file as base64 and send it back
        with open(completed_file_path, "rb") as f:
            completed_file_base64_encoded = base64.b64encode(f.read()).decode('utf-8')

        return jsonify({
            "completed_filename": completed_filename,
            "completed_file_data": completed_file_base64_encoded
        })

    except Exception as e:
        print(f"ERROR: An error occurred during file summarization: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        if temp_uploaded_filepath and os.path.exists(temp_uploaded_filepath):
            os.remove(temp_uploaded_filepath)
            print(f"DEBUG: Cleaned up temporary uploaded file: {temp_uploaded_filepath}")

# This block is for direct execution of analyze.py for testing.
# In your setup, main.py might import and run this app in a thread, so this __name__ == "__main__"
# block won't execute when run via main.py.
if __name__ == "__main__":
    app.run(port=5003, debug=True) # Running on a different port (5003) to avoid conflict with 5002
