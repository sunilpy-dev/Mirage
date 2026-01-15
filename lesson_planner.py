# planner.py

import os
import uuid
import google.generativeai as genai
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
import docx.oxml.shared
import docx.text.paragraph
from werkzeug.utils import secure_filename
import json
import base64
from apikey import GEN_AI_API_KEY # Ensure you have your API key set up in apikeys.py

# For PDF extraction
try:
    from PyPDF2 import PdfReader
except ImportError:
    print("PyPDF2 not found. Please install it: pip install PyPDF2")
    PdfReader = None

# For Image (OCR) extraction
try:
    from PIL import Image
    import pytesseract
    # Set the path to the Tesseract executable if it's not in your PATH
    # IMPORTANT: Adjust this path as needed for your Tesseract installation
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' 
except ImportError:
    print("Pillow or pytesseract not found. Please install them: pip install Pillow pytesseract")
    Image = None
    pytesseract = None
except Exception as e:
    print(f"Error importing Tesseract or Pillow: {e}. OCR functionality may be limited.")
    Image = None
    pytesseract = None


app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "www/uploaded_files"
PLANS_FOLDER = "www/lesson_plans" # New folder for lesson plans
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PLANS_FOLDER, exist_ok=True)

genai.configure(api_key=GEN_AI_API_KEY)
model = genai.GenerativeModel("gemini-2.5-flash")

def extract_text_from_docx(file_path):
    """Extracts text from a DOCX file."""
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"ERROR: Could not extract text from DOCX: {e}")
        return None

def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file."""
    if PdfReader is None:
        print("ERROR: PyPDF2 is not installed. Cannot extract text from PDF.")
        return None
    try:
        text = ""
        with open(file_path, 'rb') as f: # Use 'with open' for safer file handling
            reader = PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or "" # Handle potential None for empty pages
        return text
    except Exception as e:
        print(f"ERROR: Could not extract text from PDF: {e}")
        return None

def extract_text_from_image(file_path):
    """Extracts text from an image file using OCR."""
    if Image is None or pytesseract is None:
        print("ERROR: Pillow or pytesseract is not installed/configured. Cannot extract text from image.")
        return None
    try:
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        return text
    except pytesseract.TesseractNotFoundError: # Specific error for Tesseract not found
        print("ERROR: Tesseract OCR is not installed or not found in PATH. Image text extraction will fail.")
        return None
    except Exception as e:
        print(f"ERROR: Could not extract text from image using OCR: {e}")
        return None

# This function is not directly used in the current /plan-lesson flow for inserting content,
# but kept for consistency with file_completion.py if needed for future enhancements.
def insert_paragraph_after(doc_obj, paragraph_before, text_to_insert, style_name=None):
    """
    Inserts a new paragraph with 'text_to_insert' after 'paragraph_before'.
    'doc_obj' is the Document object.
    'style_name' is the name of a PARAGRAPH style (e.g., 'Normal', 'Heading 1').
    """
    if paragraph_before and paragraph_before.runs:
        parent = paragraph_before._element.getparent()
        new_p = docx.oxml.shared.OxmlElement("w:p")
        paragraph_before._element.addnext(new_p)
        new_paragraph = docx.text.paragraph.Paragraph(new_p, paragraph_before.document)
        new_paragraph.add_run(text_to_insert)
        if style_name:
            try:
                new_paragraph.style = doc_obj.styles[style_name]
            except KeyError:
                print(f"WARNING: Style '{style_name}' not found. Using 'Normal'.")
                new_paragraph.style = doc_obj.styles['Normal']
        return new_paragraph
    else:
        return doc_obj.add_paragraph(text_to_insert, style=style_name)

@app.route("/", methods=["GET"])
def home():
    """A simple root route to confirm the Flask service is running."""
    return "Flask lesson planner service is running!", 200

@app.route("/plan-lesson", methods=["POST"])
def plan_lesson_endpoint():
    print("DEBUG: /plan-lesson endpoint hit.")
    if not request.json:
        print("ERROR: Request must be JSON.")
        return jsonify({"error": "Request must be JSON"}), 400

    filename = request.json.get("filename")
    base64_file_data = request.json.get("file_data")
    teacher_notes = request.json.get("teacher_notes", "") # Optional notes from the teacher
    mime_type = request.json.get("mime_type") # NEW: Get mime_type for image handling

    if not filename or not base64_file_data or not mime_type: # UPDATED: Check for mime_type
        print("ERROR: Missing filename, file_data, or mime_type in request.")
        return jsonify({"error": "Missing filename, file_data, or mime_type"}), 400

    temp_uploaded_filepath = None
    extracted_text = None
    gemini_input_parts = [] # NEW: List to hold parts for Gemini model input

    try:
        # Decode the base64 data and save it as a temporary file
        file_bytes = base64.b64decode(base64_file_data)
        unique_filename = f"{uuid.uuid4()}_{secure_filename(filename)}"
        temp_uploaded_filepath = os.path.join(UPLOAD_FOLDER, unique_filename)

        with open(temp_uploaded_filepath, "wb") as f:
            f.write(file_bytes)
        print(f"DEBUG: Temporary file saved at: {temp_uploaded_filepath}")

        # Determine file type and extract text, then build gemini_input_parts
        file_extension = os.path.splitext(filename)[1].lower()
        if file_extension == ".docx":
            extracted_text = extract_text_from_docx(temp_uploaded_filepath)
            if not extracted_text or not extracted_text.strip():
                return jsonify({"error": "Could not extract meaningful text from the DOCX file. Please ensure the file contains readable content."}), 400
            gemini_input_parts.append({"text": extracted_text})
            print("DEBUG: Extracted text from DOCX.")
        elif file_extension == ".pdf":
            extracted_text = extract_text_from_pdf(temp_uploaded_filepath)
            if not extracted_text or not extracted_text.strip():
                return jsonify({"error": "Could not extract meaningful text from the PDF file. Please ensure the file contains readable content."}), 400
            gemini_input_parts.append({"text": extracted_text})
            print("DEBUG: Extracted text from PDF.")
        elif mime_type.startswith("image/"): # Use mime_type for more reliable image detection
            extracted_text = extract_text_from_image(temp_uploaded_filepath)
            
            # Add the image as inline_data to Gemini input
            gemini_input_parts.append({
                "inline_data": {
                    "mime_type": mime_type,
                    "data": base64_file_data
                }
            })
            # Add a text part to guide Gemini, including OCR'd text or a note if OCR failed
            gemini_input_parts.append({
                "text": f"Generate a lesson plan based on the content of this image, which appears to be a document or a visual aid. The extracted text from the image (if any) is:\n\n{extracted_text if extracted_text else 'No legible text was extracted by OCR. Please analyze the image content directly for content.'}"
            })
            print(f"DEBUG: Processed image and extracted text (length: {len(extracted_text) if extracted_text is not None else 0}).")
        else:
            return jsonify({"error": "Unsupported file type. Please upload DOCX, PDF, or image files."}), 400
        
        # Construct the main prompt template for the lesson plan
        lesson_plan_prompt_template = (
            f"You are an AI assistant specialized in creating detailed lesson plans for teachers. "
            f"Based on the provided content (which might include text from documents/PDFs or content from images), "
            f"generate a comprehensive lesson plan. "
            f"The lesson plan should include sections like: 'Lesson Title', 'Grade Level/Subject', 'Learning Objectives', "
            f"'Materials', 'Procedure (Step-by-Step)', 'Differentiation/Accommodations', 'Assessment', and 'Extension Activities'. "
            f"Format the output as a JSON object with these sections as keys. If a section is not applicable or cannot be generated "
            f"from the provided content, state 'N/A' or 'Not applicable'.\n\n"
            f"Ensure the JSON is well-formed and does not contain any additional text or markdown outside the JSON block. "
        )
        gemini_input_parts.append({"text": lesson_plan_prompt_template})

        if teacher_notes:
            gemini_input_parts.append({"text": f"Additional Teacher Notes:\n{teacher_notes}\n\n"})

        print(f"DEBUG: Sending prompt to Generative Model. Number of parts: {len(gemini_input_parts)}")

        response = model.generate_content(gemini_input_parts) # UPDATED: Pass list of parts
        gemini_response_text = response.text
        print("DEBUG: Received response from Generative Model.")
        print(f"DEBUG: Gemini Response: {gemini_response_text[:500]}...") # Print first 500 chars

        # Determine the output filename
        name, _ = os.path.splitext(filename)
        lesson_plan_filename = f"{name}_lesson_plan.docx"
        lesson_plan_filepath = os.path.join(PLANS_FOLDER, lesson_plan_filename)

        # Create a new DOCX document for the lesson plan
        doc = Document()
        doc.add_heading('Lesson Plan', level=1)

        # Attempt to parse Gemini's response as JSON for structured insertion
        try:
            # Clean JSON response from Gemini (remove ```json and ``` if present)
            if gemini_response_text.startswith("```json"):
                gemini_response_text = gemini_response_text[7:].strip()
            if gemini_response_text.endswith("```"):
                gemini_response_text = gemini_response_text[:-3].strip()

            lesson_plan_data = json.loads(gemini_response_text)
            if isinstance(lesson_plan_data, dict):
                for key, value in lesson_plan_data.items():
                    # Add section heading
                    doc.add_heading(key, level=2)
                    # Add content for the section
                    if isinstance(value, list):
                        for item in value:
                            doc.add_paragraph(str(item), style='List Bullet')
                    else:
                        for line in str(value).split("\n"):
                            if line.strip():
                                doc.add_paragraph(line, style='Normal')
            else:
                print("DEBUG: Gemini response is not a dictionary. Appending as plain text.")
                doc.add_paragraph("Generated Lesson Plan (Unstructured):", style='Heading 2')
                for line in gemini_response_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line, style='Normal')

        except json.JSONDecodeError:
            print("DEBUG: Gemini response is not valid JSON. Appending as plain text.")
            doc.add_paragraph("Generated Lesson Plan (Unstructured):", style='Heading 2')
            for line in gemini_response_text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line, style='Normal')

        doc.save(lesson_plan_filepath)
        print(f"DEBUG: Lesson plan saved at: {lesson_plan_filepath}")

        # Read the completed file as base64 and send it back
        with open(lesson_plan_filepath, "rb") as f:
            lesson_plan_base64_encoded = base64.b64encode(f.read()).decode('utf-8')

        return jsonify({
            "lesson_plan_filename": lesson_plan_filename,
            "lesson_plan_data": lesson_plan_base64_encoded
        })

    except Exception as e:
        print(f"ERROR: An error occurred during lesson plan generation: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        if temp_uploaded_filepath and os.path.exists(temp_uploaded_filepath):
            os.remove(temp_uploaded_filepath)
            print(f"DEBUG: Cleaned up temporary uploaded file: {temp_uploaded_filepath}")

if __name__ == "__main__":
    app.run(port=5010, debug=True) # Run on a different port than file_completion.py
